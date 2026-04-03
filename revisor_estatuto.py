import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def exibir_revisor(): # Nome ajustado para bater com o portal.py
    
    # --- CONFIGURAÇÃO DA API ---
    # Certifique-se que a chave GEMINI_API_KEY está nos Secrets do Streamlit
    if "GEMINI_API_KEY" not in st.secrets:
        st.error("Erro: GEMINI_API_KEY não encontrada nos Secrets.")
        return

    API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=API_KEY)
    
    # 1. FUNÇÃO PARA EXTRAIR TEXTO
    def extrair_texto_pdf(arquivo_pdf):
        try:
            leitor = PdfReader(arquivo_pdf)
            texto = ""
            for pagina in leitor.pages:
                conteudo = pagina.extract_text()
                if conteudo: texto += conteudo
            return texto
        except Exception as e:
            st.error(f"Erro ao ler PDF: {e}")
            return ""
    
    # 2. FUNÇÃO DE ANÁLISE COM GEMINI
    def analisar_estatuto(texto_estatuto):
        try:
            # Lista modelos para garantir compatibilidade
            modelos_disponiveis = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            modelo_escolhido = next((m for m in ['models/gemini-1.5-flash', 'models/gemini-1.5-pro'] if m in modelos_disponiveis), modelos_disponiveis[0])
            model = genai.GenerativeModel(modelo_escolhido)
            
            prompt = f"""
            Você é o Consultor Sênior da CORE ESSENCE. Analise o estatuto abaixo com base no MROSC (Lei 13.019/2014) e na Portaria Conjunta 33/2023.
            O parecer deve ser profissional, direto e estruturado para um documento oficial de consultoria governamental.
            
            Estruture a resposta exatamente assim:
            1. ✅ PONTOS DE CONFORMIDADE (O que está correto)
            2. ⚠️ GARGALOS E RISCOS (Base Portaria 33/2023)
            3. ❌ OMISSÕES OBRIGATÓRIAS (O que falta para aprovação)
            4. 💡 RECOMENDAÇÃO CORE ESSENCE (Sugestões estratégicas)
    
            Ao final, adicione o campo:
            [NOME DO CONSULTOR]
            Consultor Sênior - CORE ESSENCE
    
            Texto do Estatuto: {texto_estatuto[:28000]}
            """
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            st.error(f"Erro na Inteligência Artificial: {e}")
            return None
    
    # 3. FUNÇÃO PARA GERAR O ARQUIVO WORD (.DOCX)
    def gerar_word_parecer(texto_parecer):
        doc = Document()
        
        # Cabeçalho do Documento
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("PARECER TÉCNICO DE CONFORMIDADE ESTATUTÁRIA")
        run.bold = True
        run.font.size = Pt(14)
        
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run("CORE ESSENCE - Consultoria e Estratégia Governamental")
        run_sub.font.size = Pt(11)
        run_sub.italic = True
    
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        # Limpeza do texto para o Word (remove negritos do Markdown)
        texto_limpo = texto_parecer.replace("**", "").replace("###", "").replace("##", "")
        
        for linha in texto_limpo.split('\n'):
            if linha.strip():
                p = doc.add_paragraph(linha.strip())
                # Define fonte padrão para parecer oficial
                run_p = p.runs[0] if p.runs else p.add_run()
                run_p.font.name = 'Arial'
                run_p.font.size = Pt(11)
    
        # Rodapé de Identificação
        doc.add_paragraph("\n\n")
        footer = doc.add_paragraph("Documento gerado automaticamente pelo Sistema Core Essence para fins de revisão técnica.")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # --- INTERFACE STREAMLIT ---
    st.header("📑 Revisor de Estatuto 33/2023")
    st.write("Análise de conformidade para OSCs e Entidades do Terceiro Setor.")
    
    st.info("💡 **Diferencial:** Analisando conforme o MROSC e a **Portaria Conjunta 33/2023**. Gere a minuta em Word para facilitar seu trabalho de consultoria.")
    
    arquivo = st.file_uploader("Faça o upload do Estatuto em PDF", type=["pdf"])
    
    if arquivo:
        texto_extraido = extrair_texto_pdf(arquivo)
        if texto_extraido:
            st.success("✅ Texto extraído com sucesso!")
            
            if st.button("🚀 Iniciar Análise Estratégica"):
                with st.spinner("O Consultor IA está revisando as cláusulas..."):
                    resultado = analisar_estatuto(texto_extraido)
                    
                    if resultado:
                        st.markdown("---")
                        st.subheader("📋 Prévia do Parecer")
                        st.write(resultado)
                        
                        # Gera o arquivo Word
                        word_final = gerar_word_parecer(resultado)
                        
                        st.download_button(
                            label="📥 Baixar Parecer Completo (.docx)",
                            data=word_final,
                            file_name=f"Parecer_Tecnico_CoreEssence_{arquivo.name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

# Para permitir testes individuais do arquivo
if __name__
