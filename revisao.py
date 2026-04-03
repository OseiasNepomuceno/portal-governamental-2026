# Arquivo: revisao.py - Inteligência Core Essence
import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def exibir_revisao():
    # --- CONFIGURAÇÃO DA API ---
    # Certifique-se de que a chave GEMINI_API_KEY está nos Secrets do Streamlit Cloud
    try:
        API_KEY = st.secrets["GEMINI_API_KEY"]
        genai.configure(api_key=API_KEY)
    except Exception as e:
        st.error("Erro: Chave API do Gemini não encontrada nos Secrets.")
        return

    # 1. FUNÇÃO PARA EXTRAIR TEXTO
    def extrair_texto_pdf(arquivo_pdf):
        try:
            leitor = PdfReader(arquivo_pdf)
            texto = ""
            for pagina in leitor.pages:
                conteudo = pagina.extract_text()
                if conteudo: texto += conteudo
            return texto
        except Exception as e:
            st.error(f"Erro ao ler PDF: {e}")
            return ""

    # 2. FUNÇÃO DE ANÁLISE COM GEMINI
    def analisar_estatuto(texto_estatuto):
        try:
            # Lista modelos para garantir compatibilidade
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            prompt = f"""
            Você é o Consultor Sênior da CORE ESSENCE. Analise o estatuto abaixo com base no MROSC e na Portaria Conjunta 33/2023.
            O parecer deve ser profissional, direto e estruturado para um documento oficial de consultoria governamental.
            
            Estruture a resposta rigorosamente nestes tópicos:
            1. ✅ PONTOS DE CONFORMIDADE (O que está correto)
            2. ⚠️ GARGALOS E RISCOS (Baseado na Portaria 33/2023)
            3. ❌ OMISSÕES OBRIGATÓRIAS (O que falta para aprovação de recursos)
            4. 💡 RECOMENDAÇÃO CORE ESSENCE (Sugestão estratégica final)
    
            Ao final, adicione o campo:
            [NOME DO CONSULTOR]
            Consultor Sênior - CORE ESSENCE
    
            Texto do Estatuto: {texto_estatuto[:20000]} 
            """
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            st.error(f"Erro na Inteligência Artificial: {e}")
            return None

    # 3. FUNÇÃO PARA GERAR O ARQUIVO WORD (.DOCX)
    def gerar_word_parecer(texto_parecer):
        doc = Document()
        
        # Cabeçalho formatado
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("PARECER TÉCNICO DE CONFORMIDADE ESTATUTÁRIA")
        run.bold = True
        run.font.size = Pt(14)
        
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run("CORE ESSENCE - Consultoria e Estratégia Governamental")
        run_sub.font.size = Pt(11)
        run_sub.italic = True
    
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        # Corpo do texto limpo
        texto_limpo = texto_parecer.replace("**", "")
        
        for linha in texto_limpo.split('\n'):
            if linha.strip():
                p = doc.add_paragraph(linha)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
    
        # Rodapé
        doc.add_paragraph("\n")
        footer = doc.add_paragraph("Documento gerado automaticamente para revisão técnica Core Essence.")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- INTERFACE ---
    st.title("📜 Revisor de Estatuto (Portaria 33/2023)")
    st.markdown("---")
    
    st.info("Este módulo utiliza IA para verificar se o estatuto da entidade está apto a receber recursos conforme a nova legislação.")
    
    arquivo = st.file_uploader("Faça o upload do Estatuto em PDF", type=["pdf"])
    
    if arquivo:
        st.success("Documento carregado com sucesso!")
        if st.button("🚀 Iniciar Análise Estratégica"):
            with st.spinner("Nossa IA está processando o documento e aplicando as regras da Portaria 33..."):
                texto_extraido = extrair_texto_pdf(arquivo)
                if texto_extraido:
                    resultado = analisar_estatuto(texto_extraido)
                    
                    if resultado:
                        st.markdown("### 📋 Prévia do Parecer")
                        st.write(resultado)
                        
                        word_final = gerar_word_parecer(resultado)
                        
                        st.download_button(
                            label="📥 Baixar Parecer em Word (.docx)",
                            data=word_final,
                            file_name="Parecer_Tecnico_CoreEssence.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

# Se rodar direto para teste
if __name__ == "__main__":
    exibir_revisao()
